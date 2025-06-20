import docx
from pathlib import Path

doc_path = Path('tmp') / 'QUESTIONNAIRE EXPORT COMPTA NEW.docx'
# Use absolute path to avoid issues
doc_path = doc_path.resolve()
print('Using path:', doc_path)
print('Exists?', doc_path.exists())
print('Opening doc...')
doc = docx.Document(str(doc_path))
print('Total paragraphs:', len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs[:100]):
    text = para.text.strip()
    if text:
        print(f"{i}: {text}") 